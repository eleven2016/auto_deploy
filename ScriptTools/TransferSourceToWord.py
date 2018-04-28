
from docx import Document
import os
from ScriptTools.RecursiveReadDir import find_java_file, read_java_file_content

if __name__ == '__main__':

    '''源代码目录'''
    source_file_path = "D:\\源码学习\\spring-framework-master"
    '''源代码目录中的目标目录前缀'''
    source_file_filter = "spring-"

    '''目标目录'''
    target_word_file_path = "D:\\workspace_test\\AutoDeploy\\targetfiles"

    '''目标文件名称'''
    target_word_file_name = "spring-aop.docx"

    '''源码标题level'''
    source_title_level = 0

    '''源码模块level'''
    source_module_level = 2
    '''目录'''
    source_catalog_level = 3

    '''源码类level'''
    source_class_level = 4

    '''内容样式'''
    source_class_content_style = "No Spacing"


    '''读取源码目录,生成word文档'''
    #获取源码目录列表
    source_module_dicts = {}
    source_dirs = os.listdir(source_file_path)

    # 过滤出符合条件的目录作为 module
    for dir in source_dirs:
        if dir.startswith(source_file_filter) and not dir.endswith("spring-framework-bom"):
            module_path = source_file_path+"\\"+dir+"\\src\\main\\java\\org"
            source_module_dicts[dir] = module_path


    '''添加标题'''
    # target_word_file_document.add_heading('Spring', source_title_level)
    # 读取module下的java文件
    for module_dic in source_module_dicts:
        try:
            '''创建文件对象'''
            target_word_file_document = Document()
            target_word_file_document.add_heading(module_dic, source_module_level)
            dic_path = source_module_dicts.get(module_dic)
            print("开始生成-->"+module_dic)
            # 读取源码目录下所有的java文件
            source_class_dicts = {}
            for class_dir in os.listdir(dic_path):
                find_java_file(os.path.join(dic_path, class_dir), source_class_dicts)

            # 获取到所有的java文件,进行word文件的读写
            catalog_List = []
            for class_dic in source_class_dicts:
                class_path = source_class_dicts[class_dic]
                catalog_path = class_path[class_path.find("org")+4: class_path.find(class_dic)]
                catalog_path_temp = catalog_path.replace("\\", "/")

                if catalog_List.count(catalog_path_temp) == 0:
                    # print(catalog_path_temp)
                    catalog_List.append(catalog_path_temp)
                    target_word_file_document.add_heading(catalog_path_temp, source_catalog_level)

                target_word_file_document.add_heading(class_dic, source_class_level)
                ## 读取文件内容
                read_java_file_content(class_path, target_word_file_document)
            print("保存文件-->" + module_dic)
            '''保存文件'''
            target_word_file_document.save(target_word_file_path + "\\" + module_dic+".docx")
        except Exception as e:
            print(e)
        else:
            print("生成文件成功")





