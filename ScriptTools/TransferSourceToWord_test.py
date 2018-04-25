
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
if __name__ == '__main__':

    '''源代码目录'''
    source_file_path = "D:\\源码学习\\spring-framework-master"
    '''源代码目录中的目标目录前缀'''
    source_file_filter = "spring-"

    '''目标目录'''
    target_word_file_path = "D:\\workspace_test\\AutoDeploy\\targetfiles"

    '''目标文件名称'''
    target_word_file_name = "spring_source.docx"

    '''源码标题level'''
    source_title_level = 0

    '''源码模块level'''
    source_module_level = 2

    '''源码类level'''
    source_class_level = 4

    '''内容样式'''
    source_class_content_style = ""

    target_word_file_document = Document()
    styles = target_word_file_document.styles
    paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]

    # target_word_file_document.add_heading('Spring', source_title_level)
    # target_word_file_document.add_heading('Spring Module', source_module_level)
    # target_word_file_document.add_heading('Spring Class', source_class_level)

    for style in paragraph_styles:
        print(style.name)
        # target_word_file_document.add_paragraph("Spring source content----"+style.name, 'No Spacing')

    # target_word_file_document.add_paragraph('Spring source content2')

    # target_word_file_document.save(target_word_file_path+"\\"+target_word_file_name)


